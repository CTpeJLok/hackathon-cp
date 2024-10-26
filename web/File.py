import tempfile
import io

import openpyxl
import xlrd
from docx import Document
import pdfplumber


# Базовый класс для обработки файлов
class File:
    def __init__(self, file_name: str, file_bytes: bytes):
        self.file_name = file_name  # Имя файла
        self.file_bytes = io.BytesIO(file_bytes)  # Байтовый поток с содержимым файла

    def get_table(self):
        # Возвращаем байтовый поток (может быть переопределено в подклассах)
        return self.file_bytes


# Класс для работы с файлами формата .xlsx
class XLSXFile(File):
    def __init__(self, file_name: str, file_bytes: bytes):
        super().__init__(file_name, file_bytes)  # Инициализация родительского класса

    def get_table(self):
        # Чтение файла Excel формата .xlsx
        wb_x: openpyxl.Workbook = openpyxl.load_workbook(
            self.file_bytes, data_only=True
        )
        sheet_x = wb_x.active  # Получаем активный лист
        if not sheet_x:
            return []  # Возвращаем пустой список, если листа нет

        table = []
        for row in sheet_x.iter_rows(values_only=True):
            table.append(list(row))  # Добавляем строки в таблицу

        return table  # Возвращаем таблицу


# Класс для работы с файлами формата .xls
class XLSFile(File):
    def __init__(self, file_name: str, file_bytes: bytes):
        super().__init__(file_name, file_bytes)

    def get_table(self):
        # Используем временный файл для хранения содержимого .xls
        with tempfile.NamedTemporaryFile(delete=False, suffix=".xls") as temp_file:
            temp_file.write(self.file_bytes.getbuffer())
            temp_file_path = temp_file.name  # Получаем путь к временному файлу

            # Чтение файла Excel формата .xls
            wb: xlrd.book.Book = xlrd.open_workbook(temp_file_path)
            sheet: xlrd.sheet.Sheet = wb.sheet_by_index(0)  # Получаем первый лист

            table = []
            for r in range(sheet.nrows):
                row_data = []  # Список для строки
                for c in range(sheet.ncols):
                    row_data.append(sheet.cell_value(r, c))  # Добавляем ячейки в строку
                table.append(row_data)  # Добавляем строку в таблицу

            return table  # Возвращаем таблицу


# Класс для работы с файлами формата .docx
class DOCXFile(File):
    def __init__(self, file_name: str, file_bytes: bytes):
        super().__init__(file_name, file_bytes)

    def get_table(self):
        # Чтение таблицы из файла формата .docx
        doc = Document(self.file_bytes)  # Открываем документ

        table = []

        doc_table = doc.tables[:1]  # Получаем первую таблицу
        if not doc_table:
            return []  # Возвращаем пустой список, если таблицы нет

        doc_table = doc.tables[0]  # Сохраняем первую таблицу
        for r in doc_table.rows:
            row_data = []  # Список для строки
            for c in r.cells:
                row_data.append(c.text)  # Добавляем текст ячеек в строку
            table.append(row_data)  # Добавляем строку в таблицу

        return table  # Возвращаем таблицу


# Класс для работы с файлами формата .pdf
class PDFFile(File):
    def __init__(self, file_name: str, file_bytes: bytes):
        super().__init__(file_name, file_bytes)

    def get_table(self):
        # Используем временный файл для хранения содержимого .pdf
        with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as temp_file:
            temp_file.write(self.file_bytes.getbuffer())
            temp_file_path = temp_file.name  # Получаем путь к временному файлу

            # Чтение таблицы из файла формата .pdf
            with pdfplumber.open(temp_file_path) as pdf:
                all_tables = []  # Список для всех таблиц
                for page in pdf.pages:
                    table = page.extract_table()  # Извлекаем таблицу из страницы
                    if table:
                        all_tables.append(
                            table
                        )  # Добавляем таблицу, если она существует

                tables = []  # Список для окончательных таблиц

                for t in all_tables:
                    for r in t:
                        if r:
                            tables.append(list(r))  # Добавляем строки в таблицу

                return tables  # Возвращаем таблицы
